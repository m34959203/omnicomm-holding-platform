#!/usr/bin/env python3
"""Сборка Word-документа ТЗ (docx) из содержания docs/TZ-holding.md.

Кладёт результат в web/public/ и web/out/ — чтобы файл раздавался реверс-прокси
по публичной ссылке https://omnicomm.technokod.kz/TZ-Omnicomm-Holding.docx
(web/public сохраняется при ребилде, web/out — текущая раздача).

Запуск: python3 scripts/build_tz_docx.py
"""

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "ТЗ-Omnicomm-Holding.docx"
PUBLIC = ROOT / "web" / "public" / "TZ-Omnicomm-Holding.docx"
SERVED = ROOT / "web" / "out" / "TZ-Omnicomm-Holding.docx"

INK = RGBColor(0x1A, 0x1A, 0x1A)
SLATE = RGBColor(0x1F, 0x3A, 0x3D)
ACCENT = RGBColor(0x2E, 0x7D, 0x6B)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
FONT = "Calibri"


def _f(run, name=FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        from docx.oxml import OxmlElement
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), name)


def base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT; st.font.size = Pt(10.5); st.font.color.rgb = INK
    st.paragraph_format.space_after = Pt(5); st.paragraph_format.line_spacing = 1.12


def h1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = SLATE; _f(r)


def body(doc, text, bold=False, color=INK, size=10.5, after=5):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color; _f(r)
    return p


def req(doc, code, text, status=""):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    if code:
        r = p.add_run(code + "  "); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = ACCENT; _f(r)
    r = p.add_run(text); r.font.size = Pt(10.5); _f(r)
    if status:
        r = p.add_run("  " + status); r.font.size = Pt(10.5); r.bold = True; _f(r)


def shade(cell, hexc):
    from docx.oxml import OxmlElement
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(sh)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10); _f(r); shade(c, "2E7D6B")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; r = cells[i].paragraphs[0].add_run(v); r.font.size = Pt(9.5); _f(r)
    return t


def build():
    doc = Document(); base(doc)
    s = doc.sections[0]; s.left_margin = s.right_margin = Pt(60)

    # титул
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ТЕХНИЧЕСКОЕ ЗАДАНИЕ"); r.font.size = Pt(13); r.bold = True; r.font.color.rgb = ACCENT; _f(r)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Аналитическая платформа автопарка холдинга"); r.font.size = Pt(24); r.bold = True
    r.font.color.rgb = SLATE; _f(r)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("на базе телематики Omnicomm"); r.font.size = Pt(15); r.font.color.rgb = MUTED; _f(r)
    doc.add_paragraph()
    for k, v in [("Заказчик", "АО НАК «Казатомпром» (23 ДЗО, ~2000 ТС)"),
                 ("Исполнитель", "TechnoKod"), ("Версия", "1.0"), ("Дата", "22 июня 2026")]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{k}: "); r.bold = True; r.font.color.rgb = MUTED; r.font.size = Pt(10.5); _f(r)
        r = p.add_run(v); r.font.size = Pt(10.5); _f(r)
    doc.add_page_break()

    body(doc, "Обозначения: ✅ реализовано · 🟡 частично · ❌ не реализовано · 🔵 MVP-приоритет заказчика.",
         color=MUTED, size=9.5)

    h1(doc, "0. Паспорт")
    table(doc, ["Параметр", "Значение"], [
        ["Продукт", "Аналитическая платформа автопарка холдинга поверх Omnicomm Online"],
        ["Заказчик", "АО НАК «Казатомпром» (КАП), 23 ДЗО, ~2000 ТС"],
        ["Источник данных", "Omnicomm Online REST API (service omnicomm)"],
        ["Бизнес-модель", "Аналитика как услуга, без продажи железа; раздача по ДЗО без лицензий"],
        ["Инфраструктура", "В РК, без трансграничной передачи ПДн"],
    ])

    h1(doc, "1. Роли и контроль доступа")
    req(doc, "", "Две роли без тонкого RBAC: «начальник» (весь холдинг) и «ДЗО» (своё поддерево dim_org).")
    req(doc, "", "Row-level доступ по поддереву — жёсткое требование: ДЗО видит свою технику + под-ДЗО + подрядчиков, не видит соседей.")
    req(doc, "", "Одинаковый дашборд для всех (просмотр + экспорт), различие только в видимом scope. Без редактирования данных.")
    req(doc, "", "Multi-tenant: у каждого ДЗО свой логин/пароль.", "✅")

    h1(doc, "2. Сбор и нормализация данных")
    req(doc, "R2.1", "Забор телеметрии из Omnicomm REST: дерево ТС, сводный отчёт (POST), геозоны, трек, активность.", "✅")
    req(doc, "R2.2", "Иерархия организаций dim_org (Холдинг → ДЗО → под-ДЗО → подрядчик) из дерева ТС.", "✅ (73 узла)")
    req(doc, "R2.3", "Привязка ТС → организация по terminal_id; роллапы KPI по поддереву.", "✅")
    req(doc, "R2.4", "Настраиваемая частота обновления данных.", "🔵 🟡")
    req(doc, "R2.5", "Единицы Omnicomm: топливо дл→л, время сек→ч; чистка GPS-выбросов по физике.", "✅")

    h1(doc, "3. Дашборды и отчёты")
    req(doc, "R3.1", "Executive-дашборд холдинга: топлайн (топливо/потери ₸), плитки ДЗО, drill в узел.", "✅")
    req(doc, "R3.2", "Дашборд на ДЗО: KPI, «куда уходят деньги», первоочередные ТС, телеметрия, под-организации.", "✅/🟡")
    req(doc, "R3.3", "Выгрузка в Excel (главный приоритет) — кнопка в портале на текущий срез.", "🔵 🟡")
    req(doc, "R3.4", "Два языка (RU / KK) — i18n портала и отчётов.", "🔵 ❌")
    req(doc, "R3.5", "Мобильная версия (просмотр с телефона).", "🔵 🟡")
    req(doc, "R3.6", "Отчёт .pptx/.html(+PDF) уровня презентации.", "✅")

    h1(doc, "4. Скоростной режим и нарушения")
    req(doc, "R4.1", "Модель geozone_limit(геозона, тип ТС) по СТ КАП: ~400 именованных геозон + матрица 6 зон × 3 категории (фолбэк).", "❌")
    req(doc, "R4.2", "Детекция превышений по GPS-треку против лимита; классификация по диапазонам.", "❌")
    req(doc, "R4.3", "Правило min(внутренний, ПДД); пороги СТ КАП ≤3 / 3–5 / ≥6 (грубое).", "❌")
    req(doc, "R4.4", "Спец-лимиты по типу груза/ТС (опасные грузы −30 км/ч и т.п.).", "❌")

    h1(doc, "5. Рекомендации «на букве закона»")
    req(doc, "R5.1", "Каждая рекомендация: нарушение → статья КоАП РК → частота → сумма штрафов (МРП 2026 = 4325 ₸) → ущерб → действие.", "❌")
    req(doc, "R5.2", "Справочник статей/штрафов в конфиге (ст. 592/571, МРП).", "❌")
    req(doc, "R5.3", "Досье водителя при наличии идентификации (при правовом основании).", "❌")
    req(doc, "R5.4", "Лестница эскалации: разовое → беседа → приказ → дооснащение DSM/ADAS (upsell).", "❌")
    req(doc, "R5.5", "Бизнес-инвариант: без обвинений, «требует проверки»; строго по статье.", "✅")

    h1(doc, "6. Контроль ТО")
    req(doc, "R6.1", "Авто-учёт наработки «от нуля» (T0 = подключение) по моточасам/пробегу.", "❌")
    req(doc, "R6.2", "Интервал ТО по категории/модели (справочник).", "❌")
    req(doc, "R6.3", "Заблаговременный алерт «осталось N моточасов / M км».", "❌")
    req(doc, "R6.4", "Цикл подтверждения: ответственный подтверждает «ТО пройдено» → сброс счётчика.", "❌")
    req(doc, "R6.5", "Не дублировать Omnicomm-ТО.", "❌")

    h1(doc, "7. Качество данных (Sensor Health)")
    req(doc, "R7.1", "Терминальный «светофор» (давность данных) по activity/vehicles + receive_data.", "❌ (API доступно)")
    req(doc, "R7.2", "Дата-валидация: пометка ТС без топлива/CAN/оборотов → исключение из KPI.", "❌")
    req(doc, "R7.3", "Реестр датчиков с серийниками и хронологией замен (из монтажных отчётов).", "❌")
    req(doc, "R7.4", "Сенсор-уровневая детекция (по «Журналу») — только при доступе к API.", "⛔ заблокировано")

    h1(doc, "8. Нефункциональные требования")
    for t in [
        "NF1 Безопасность/ПДн: креды только из ENV; данные в РК; пароли клиентов — шифрование (не base64 в проде).",
        "NF2 Доступ: fail-closed изоляция ДЗО (перепроверка scope перед рендером).",
        "NF3 Масштаб: ~2000 ТС, 23 ДЗО; роллапы без деградации; SQLite сейчас, Postgres при росте.",
        "NF4 Надёжность сбора: батчинг ≤50 ТС / ≤31 дня, ретраи, упреждающий refresh JWT, коды ошибок.",
        "NF5 Локализация: RU/KK; денежные величины в ₸ (млн/тыс).",
        "NF6 Доставка: веб-портал + Excel; мобильный просмотр.",
    ]:
        req(doc, "", t)

    h1(doc, "9. MVP vs Дальше")
    body(doc, "MVP-доставка (приоритет директора):", bold=True, after=2)
    body(doc, "Excel-экспорт → два языка RU/KK → мобильная версия → логин на ДЗО [есть] → настраиваемая частота. "
              "База: иерархия + дашборды ДЗО + экономика (есть).")
    body(doc, "MVP-доказательство ценности:", bold=True, after=2)
    body(doc, "хотя бы один срез уникальной ценности на реальных данных КАП — скоростной режим по СТ КАП → "
              "дисциплинарка/штраф; плюс терминальный Sensor Health как доверие к цифрам.")
    body(doc, "Дальше (ядро ценности):", bold=True, after=2)
    body(doc, "полный geozone_limit + детекция превышений → движок рекомендаций на букве закона → контроль ТО → реестр датчиков.")
    body(doc, "До прод-онбординга (P0):", bold=True, after=2)
    body(doc, "криптостойкое хранение паролей; процесс онбординга ДЗО + приёмочный тест изоляции; Business Case (плательщик/цена/ROI).")

    h1(doc, "10. Критерии приёмки (выборка)")
    for t in [
        "ДЗО входит под своим логином и видит только свою технику; начальник — весь КАП.",
        "Дашборд ДЗО выгружается в Excel одной кнопкой; интерфейс на RU и KK; читается с телефона.",
        "Превышение квалифицируется по СТ КАП (лимит геозоны × тип ТС) и по КоАП РК с суммой штрафа.",
        "ТС с битым датчиком топлива помечен и не искажает топливные рейтинги.",
        "Контроль ТО: наработка из фида, алерт срабатывает, подтверждение сбрасывает цикл.",
    ]:
        req(doc, "", t)

    h1(doc, "11. Ограничения и допущения")
    for t in [
        "Видео/ADAS/DSM — отдельный контур интегратора, в Omnicomm не передаётся.",
        "«Журнал» (сырьё по датчикам) через REST недоступен → Sensor Health пока терминального уровня.",
        "Серийники датчиков в фиде не приходят → реестр ведётся вручную из монтажных отчётов.",
        "Контроль ТО «от нуля» — точность растёт по мере накопления; нужен T0-offset при онбординге.",
        "Справочник геозон — только verified-артефакт (ручное ревью заказчика), не авто-парс.",
        "Досье водителя — только при правовом основании, минимизация полей, без авто-решений по возрасту.",
        "online.omnicomm.ru — РФ-домен → подтвердить геолокацию копии/каналов перед заявлением «без трансграничной ПДн».",
        "Профильные системы учёта Казатомпрома недоступны.",
    ]:
        req(doc, "", t)

    doc.save(str(OUT))
    PUBLIC.parent.mkdir(parents=True, exist_ok=True)
    SERVED.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUT, PUBLIC)
    shutil.copy2(OUT, SERVED)
    print("saved:", OUT, OUT.stat().st_size, "bytes")
    print("served:", SERVED)


if __name__ == "__main__":
    build()
