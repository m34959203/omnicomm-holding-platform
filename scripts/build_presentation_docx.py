#!/usr/bin/env python3
"""Сборка клиентской ПРЕЗЕНТАЦИИ платформы (docx) — для показа заказчику КАП.

Документ-презентация (книжная ориентация): каждый раздел = «слайд» с кикер-номером,
крупным заголовком и тезисами. Фирменный teal-стиль (как Концепция). Без скриншотов.
Запуск: python3 scripts/build_presentation_docx.py
Результат: docs/Презентация-Omnicomm-Holding.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "docs" / "Презентация-Omnicomm-Holding.docx"

INK = RGBColor(0x1A, 0x1A, 0x1A)
SLATE = RGBColor(0x1F, 0x3A, 0x3D)
ACCENT = RGBColor(0x2E, 0x7D, 0x6B)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BASE_FONT = "Calibri"


def set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = BASE_FONT
    st.font.size = Pt(11)
    st.font.color.rgb = INK
    st._element.rPr.rFonts.set(qn("w:cs"), BASE_FONT)
    pf = st.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def _runfont(run, name=BASE_FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), name)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def body(doc, text, bold=False, color=INK, size=11, after=6, before=0, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    _runfont(r)
    return p


def bullet(doc, runs, size=11):
    p = doc.add_paragraph(style="List Bullet")
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        _runfont(r)
    return p


def slide(doc, number, kicker, title, first=False):
    """Начало нового «слайда»: разрыв страницы + кикер-номер + крупный заголовок."""
    if not first:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{number:02d} · {kicker}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT
    _runfont(r)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(23)
    r.font.color.rgb = SLATE
    _runfont(r)
    # тонкая акцентная линия под заголовком
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "2E7D6B")
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        rp = c.paragraphs[0].add_run(htext)
        rp.bold = True
        rp.font.color.rgb = WHITE
        rp.font.size = Pt(10.5)
        _runfont(rp)
        shade(c, "2E7D6B")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rp = cells[i].paragraphs[0].add_run(val)
            rp.font.size = Pt(10)
            if i == 0:
                rp.bold = True
                rp.font.color.rgb = SLATE
            _runfont(rp)
        if ri % 2 == 1:
            for c in cells:
                shade(c, "F1F6F4")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
    return t


def cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ПРЕЗЕНТАЦИЯ РЕШЕНИЯ"); r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = ACCENT; _runfont(r)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Аналитическая платформа\nавтопарка холдинга")
    r.font.size = Pt(30); r.bold = True; r.font.color.rgb = SLATE; _runfont(r)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("на базе телематики Omnicomm")
    r.font.size = Pt(16); r.font.color.rgb = MUTED; _runfont(r)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Кросс-ДЗО аналитика · экономика в ₸ · скоростной режим на букве закона")
    r.font.size = Pt(11); r.italic = True; r.font.color.rgb = ACCENT; _runfont(r)
    for _ in range(3):
        doc.add_paragraph()
    meta = [
        ("Заказчик", "АО НАК «Казатомпром» (КАП)"),
        ("Исполнитель", "TechnoKod"),
        ("Версия", "1.0"),
        ("Дата", "22 июня 2026"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in meta:
        cells = t.add_row().cells
        rp = cells[0].paragraphs[0].add_run(k); rp.bold = True
        rp.font.color.rgb = MUTED; rp.font.size = Pt(10.5); _runfont(rp)
        rp = cells[1].paragraphs[0].add_run(v)
        rp.font.size = Pt(10.5); _runfont(rp)
        cells[0].width = Pt(120); cells[1].width = Pt(260)


def build():
    doc = Document()
    set_base_style(doc)
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Pt(64)
    sec.top_margin = Pt(56); sec.bottom_margin = Pt(56)

    cover(doc)

    # 01 — О чём система
    slide(doc, 1, "СУТЬ", "Единый «мозг автопарка» поверх Omnicomm", first=False)
    body(doc, "Платформа собирает телеметрию по всем ДЗО холдинга через API Omnicomm, "
              "нормализует её в иерархию организаций (Холдинг → ДЗО → под-ДЗО → подрядчик), "
              "переводит в деньги (₸) и раздаёт каждому ДЗО свой дашборд и отчёты — без "
              "посемест­ных лицензий и ручных выгрузок, с доступом строго к своей технике.",
              size=12, after=10)
    bullet(doc, [("~2000 единиц техники", True), (" в ", False), ("23 ДЗО", True),
                 (" — единая картина холдинга", False)])
    bullet(doc, [("Не заменяет Omnicomm", True), (" — это аналитический слой поверх него", False)])
    bullet(doc, [("Главное отличие", True), (" — переводит телеметрию в деньги, закон и конкретные действия", False)])

    # 02 — Проблема
    slide(doc, 2, "ПРОБЛЕМА", "Что не закрывает текущая витрина")
    body(doc, "Сегодня управленческая аналитика — Power BI «Мониторинг скоростного "
              "режима» (~4 страницы, только скорость). Её ограничения:", after=8)
    for t, rest in [
        ("Только скорость", " — нет топливной и экономической аналитики (перерасход в ₸), нет КПД спецтехники, контроля ТО и качества данных."),
        ("Дорогая ручная доставка", " — посемест­ные лицензии Power BI (≈8500 ₸/мес × получатель) плюс ручная выгрузка каждому ДЗО раз в неделю."),
        ("Модель не масштабируется", " — «колонка на каждое ДЗО» ломается при добавлении ДЗО, под-ДЗО или подрядчика."),
        ("Рекомендации — «ИИ-вода»", " — общий план без привязки к статье закона, сумме штрафа и конкретному действию."),
    ]:
        bullet(doc, [(t, True), (rest, False)])
    body(doc, "При этом у холдинга уже есть СТ Казатомпром — стандарт безопасной эксплуатации "
              "с матрицей лимитов по зонам риска × типам ТС и ~400 именованными геозонами. "
              "Сегодня он живёт на бумаге, а не в аналитике нарушений и ущерба.", before=4, color=MUTED, size=10.5)

    # 03 — Решение / видение
    slide(doc, 3, "РЕШЕНИЕ", "Слой, которого нет ни у кого")
    body(doc, "Платформа не переизобретает то, что уже считает Omnicomm (расход, заправки/сливы, "
              "моточасы, SafeDrive) и не лезет в видео-железо. Её ценность — именно тот слой, "
              "которого нет ни у Omnicomm, ни у Power BI, ни у видеоинтегратора:", after=8)
    bullet(doc, [("Кросс-ДЗО агрегация", True), (" — сводим весь холдинг в одну иерархию и сравниваем ДЗО между собой", False)])
    bullet(doc, [("Экономика в ₸", True), (" — перерасход, простой, износ и потенциал экономии в деньгах", False)])
    bullet(doc, [("Рекомендации на букве закона", True), (" — нарушение → статья → штраф → ущерб → действие", False)])
    bullet(doc, [("Контроль ТО и качества данных", True), (" — наработка, алерты, «светофор» исправности терминалов", False)])

    # 04 — Ценность для ДЗО
    slide(doc, 4, "ЦЕННОСТЬ", "Что получает каждый ДЗО")
    add_table(doc,
        ["Пласт", "Что даёт", "Кому больно сейчас"],
        [
            ["Экономика топлива", "Перерасход/экономия в ₸, нормы, рейтинги ДЗО", "нет в Power BI"],
            ["Скоростной режим на букве закона", "Превышения по геозоне × типу ТС, квалификация по КоАП РК, сумма штрафов, ущерб, действие", "Power BI = факт, не деньги/закон"],
            ["КПД спецтехники", "Моточасы, простой под нагрузкой, ₸/моточас (буровые, погрузка)", "не считает никто"],
            ["Контроль ТО", "Авто-учёт наработки, заблаговременный алерт, подтверждение цикла", "вручную / никак"],
            ["Качество данных", "«Светофор» терминалов, пометка ТС с битыми датчиками", "аналитик ловит глазами"],
            ["Доставка", "Дашборд на ДЗО + Excel + 2 языка + мобильный вид, без лицензий", "лицензии + ручной экспорт"],
        ],
        widths=[Pt(110), Pt(250), Pt(120)])

    # 05 — Главный дифференциатор
    slide(doc, 5, "ДИФФЕРЕНЦИАТОР", "Рекомендации на букве закона")
    body(doc, "Вместо общих советов — детерминированная цепочка от факта к действию, "
              "которую можно показать руководителю и водителю:", after=10)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Нарушение  →  статья КоАП РК  →  частота  →  сумма штрафов (МРП × ₸)  "
                  "→  вероятный ущерб  →  конкретное действие")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = SLATE; _runfont(r)
    body(doc, "")
    bullet(doc, [("Штраф по КоАП — только для дорог общего пользования", True),
                 (" (на технологических дорогах — дисциплинарная мера по СТ КАП).", False)])
    bullet(doc, [("Действие соразмерно частоте", True),
                 (" — от беседы с водителем до дооснащения DSM/ADAS при систематике (инструмент продаж интегратора).", False)])
    bullet(doc, [("ИИ только переформулирует посчитанные факты", True),
                 (" в деловой текст — он не источник права и работает с graceful-fallback.", False)])

    # 06 — Скоростной режим по СТ КАП
    slide(doc, 6, "БЕЗОПАСНОСТЬ", "Скоростной режим по СТ Казатомпром")
    body(doc, "Лимит скорости считается не «вообще», а по правилу", after=2)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("лимит = min(ПДД, лимит геозоны, внутренний стандарт СТ КАП)")
    r.bold = True; r.font.size = Pt(12.5); r.font.color.rgb = ACCENT; _runfont(r)
    body(doc, "", after=4)
    bullet(doc, [("Лимит на геозону × тип ТС", True), (" — буровой нельзя 80, самосвалу на трассе можно. Этого Power BI/Omnicomm толком не умеют.", False)])
    bullet(doc, [("~400 именованных геозон СТ КАП", True), (" на интерактивной карте (Яндекс-гибрид: спутник + подписи), клик → название и лимит.", False)])
    bullet(doc, [("Чистка GPS-выбросов по физике движения", True), (" (расстояние/время между точками), а не грубым окном по длительности.", False)])
    bullet(doc, [("Честная граница", True), (" — детекция аварийности/усталости водителя требует видео-железа (DSM/ADAS) — это отдельный контур.", False)])

    # 07 — Экономика и спецтехника
    slide(doc, 7, "ДЕНЬГИ", "Экономика топлива и КПД спецтехники")
    body(doc, "Платформа показывает не литры и часы, а где уходят деньги и сколько можно вернуть:", after=8)
    bullet(doc, [("Куда уходят деньги", True), (" — ранжированные утечки: перерасход к нормам, холостой ход, простой; в каждой строке — как посчитано.", False)])
    bullet(doc, [("КПД спецтехники", True), (" — моточасы, продуктивная работа под нагрузкой против непродуктивного простоя, ₸/моточас (буровые, погрузчики).", False)])
    bullet(doc, [("Счётчик подтверждённой экономии", True), (" — честный baseline по методике IPMVP-lite: показывает фактический эффект, а не «бумажную» экономию.", False)])
    bullet(doc, [("Без обвинений", True), (" — аномалии помечаются «требует проверки», перерасход — только при согласованных нормах.", False)])

    # 08 — Доступ и доставка
    slide(doc, 8, "ДОСТУП", "Каждому ДЗО — своё, без лицензий")
    body(doc, "Снимаем главную дилемму Power BI «лицензия на каждого ИЛИ общая ссылка»:", after=8)
    bullet(doc, [("Row-level доступ", True), (" — руководство видит весь КАП, ДЗО — только своё поддерево; соседи друг друга не видят.", False)])
    bullet(doc, [("Без посемест­ных лицензий", True), (" — у каждого ДЗО свой логин, дашборд и экспорт в своём периметре.", False)])
    bullet(doc, [("MVP-приоритеты заказчика", True), (" — Excel-экспорт, два языка (RU/KK), мобильная версия, настраиваемая частота обновления.", False)])
    bullet(doc, [("Быстро", True), (" — дашборд открывается из кэш-снапшота мгновенно, синхронизация с Omnicomm идёт фоном.", False)])

    # 09 — Позиционирование
    slide(doc, 9, "ПОЗИЦИЯ", "С кем мы не конкурируем")
    for t, rest in [
        ("Omnicomm Online", " — мы надстройка, не замена. Omnicomm остаётся источником телеметрии и считает per-ТС; мы агрегируем по холдингу и переводим в решения."),
        ("Power BI", " — заменяем как витрину: дешевле (без лицензий на получателя), масштабируемо (иерархия), богаче (деньги / закон / ТО)."),
        ("Видеоинтегратор (IT-Transport)", " — у них «глаза и нервы» (DMS/ADAS/MDVR, отдельный контур). У них нет ИИ-аналитики, скоринга и холдингового среза — это наш участок."),
    ]:
        bullet(doc, [(t, True), (rest, False)])
    body(doc, "Бизнес-модель — «аналитика как услуга» поверх существующей телеметрии холдинга, "
              "без продажи железа. Вся инфраструктура — в РК, без трансграничной передачи ПДн.",
              before=6)

    # 10 — Как работает
    slide(doc, 10, "АРХИТЕКТУРА", "Как это работает — быстро и масштабируемо")
    for t, rest in [
        ("Python-движки", " — нормализация, экономика, скоростной режим СТ КАП, роллапы по иерархии (отлаженное ядро single-client отчёта)."),
        ("FastAPI-мост", " — синхронизация с Omnicomm идёт фоновой задачей с прогрессом и параллельным забором данных."),
        ("Кэш-снапшот", " — чтения дашборда отдаются из кэша мгновенно и не дёргают Omnicomm на каждый запрос: интерфейс всегда быстрый."),
        ("Next.js-интерфейс", " — современный дашборд ДЗО, интерактивная карта геозон, рекомендации."),
        ("Устойчивость", " — синк изолирован: его сбой или медлительность Omnicomm не роняет сервер; данные переживают перезапуск."),
    ]:
        bullet(doc, [(t, True), (rest, False)])

    # 11 — Статус
    slide(doc, 11, "СТАТУС", "Что уже работает (на 22.06.2026)")
    bullet(doc, [("Подключена боевая копия КАП", True), (" — дерево развёрнуто в 73 узла / ~1998 ТС; топливо и обороты идут по API.", False)])
    bullet(doc, [("Получен СТ КАП с геозонами", True), (" — ~400 зон загружены и отображаются на интерактивной карте (Яндекс-гибрид).", False)])
    bullet(doc, [("Платформа развёрнута на постоянном домене", True), (" — omnicomm.technokod.kz, доступна для демонстрации.", False)])
    bullet(doc, [("Holding-слой собран", True), (" — иерархия организаций, row-level доступ, роллапы KPI, дашборды на ДЗО.", False)])

    # 12 — Дорожная карта
    slide(doc, 12, "ДАЛЬШЕ", "Дорожная карта")
    add_table(doc,
        ["Этап", "Содержание"],
        [
            ["Ближайшее", "Движок превышений по GPS-треку × geozone_limit; рекомендации на букве закона в портале"],
            ["MVP-доставка", "Excel-экспорт по кнопке, второй язык (KK), мобильная версия"],
            ["Онбординг ДЗО", "Учётки ДЗО, тест изоляции доступа, настраиваемая частота обновления"],
            ["Развитие", "Контроль ТО, Sensor Health, расширение экономики и счётчика экономии"],
        ],
        widths=[Pt(130), Pt(350)])
    body(doc, "Порядок и сроки этапов согласуются с заказчиком.", before=6, color=MUTED, size=10)

    # 13 — Закрытие
    slide(doc, 13, "ИТОГ", "Почему это решение")
    bullet(doc, [("Дешевле", True), (" — без посемест­ных лицензий Power BI и ручных выгрузок.", False)], size=12)
    bullet(doc, [("Масштабируемо", True), (" — иерархия организаций вместо «колонки на ДЗО».", False)], size=12)
    bullet(doc, [("Богаче", True), (" — деньги, закон, ТО и качество данных там, где раньше была только скорость.", False)], size=12)
    bullet(doc, [("Безопасно", True), (" — данные в РК, доступ строго по своему ДЗО.", False)], size=12)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TechnoKod · аналитика автопарка как услуга поверх Omnicomm")
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = MUTED; _runfont(r)

    doc.save(str(OUT))
    print("saved:", OUT, OUT.stat().st_size, "bytes")


if __name__ == "__main__":
    build()
