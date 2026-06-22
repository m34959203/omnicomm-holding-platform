#!/usr/bin/env python3
"""Сборка Word-документа концепции платформы (docx) из содержания CONCEPT.md.

Формальный документ для заказчика (КАП): титул, разделы 1–8, таблицы ценности и
позиционирования. Запуск: python3 scripts/build_concept_docx.py
Результат: docs/Концепция-Omnicomm-Holding.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "docs" / "Концепция-Omnicomm-Holding.docx"

INK = RGBColor(0x1A, 0x1A, 0x1A)
SLATE = RGBColor(0x1F, 0x3A, 0x3D)      # тёмный сине-зелёный для заголовков
ACCENT = RGBColor(0x2E, 0x7D, 0x6B)     # акцент (корпоративный teal)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
BASE_FONT = "Calibri"


def set_base_style(doc: Document) -> None:
    st = doc.styles["Normal"]
    st.font.name = BASE_FONT
    st.font.size = Pt(11)
    st.font.color.rgb = INK
    st._element.rPr.rFonts.set(qn("w:cs"), BASE_FONT)
    pf = st.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def _runfont(run, name=BASE_FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)


def h1(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{number}.  {text}")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = SLATE
    _runfont(r)
    return p


def body(doc, text, bold=False, color=INK, size=11, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    _runfont(r)
    return p


def bullet(doc, runs):
    """runs: список (text, bold) для формирования строки маркера."""
    p = doc.add_paragraph(style="List Bullet")
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(11)
        _runfont(r)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        rp = c.paragraphs[0].add_run(htext)
        rp.bold = True
        rp.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        rp.font.size = Pt(10.5)
        _runfont(rp)
        shade(c, "2E7D6B")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rp = cells[i].paragraphs[0].add_run(val)
            rp.font.size = Pt(10)
            _runfont(rp)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
    return t


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("КОНЦЕПЦИЯ ПРОЕКТА"); r.font.size = Pt(13); r.font.color.rgb = ACCENT
    r.bold = True; _runfont(r)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Аналитическая платформа автопарка холдинга")
    r.font.size = Pt(26); r.bold = True; r.font.color.rgb = SLATE; _runfont(r)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("на базе телематики Omnicomm")
    r.font.size = Pt(16); r.font.color.rgb = MUTED; _runfont(r)

    for _ in range(2):
        doc.add_paragraph()

    meta = [
        ("Заказчик", "АО НАК «Казатомпром» (КАП)"),
        ("Исполнитель", "TechnoKod"),
        ("Версия", "1.1"),
        ("Дата", "22 июня 2026"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in meta:
        cells = t.add_row().cells
        rp = cells[0].paragraphs[0].add_run(k); rp.bold = True
        rp.font.color.rgb = MUTED; rp.font.size = Pt(10.5); _runfont(rp)
        rp = cells[1].paragraphs[0].add_run(v)
        rp.font.size = Pt(10.5); _runfont(rp)
        cells[0].width = Pt(120); cells[1].width = Pt(260)
    doc.add_page_break()


def build():
    doc = Document()
    set_base_style(doc)
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Pt(64)

    cover(doc)

    body(doc, "Документ верхнего уровня: зачем проект, для кого, какую ценность даёт "
              "и чем отличается. Связанные документы: Техническое задание, Техническое "
              "решение, База знаний, Целевая архитектура.", color=MUTED, size=10)

    h1(doc, 1, "Проблема")
    body(doc, "Холдинг КАП эксплуатирует ~2000 единиц техники в 23 ДЗО (рудники, логистика, "
              "спецтехника, легковой и вахтовый транспорт) на телематике Omnicomm Online. "
              "Текущая управленческая витрина — Power BI «Мониторинг скоростного режима» "
              "(~4 страницы, только скорость). Её ограничения:")
    for t, b, rest in [
        ("Только безопасность/скорость", True, " — нет топливной и экономической аналитики (₸), КПД спецтехники, контроля ТО и качества данных."),
        ("Доставка дорогая и ручная", True, " — посемест­ные лицензии Power BI (≈8500 ₸/мес × получатель) + ручная выгрузка каждому ДЗО раз в неделю."),
        ("Модель не масштабируется", True, " — «колонка на каждое ДЗО» ломается при добавлении ДЗО/под-ДЗО/подрядчика."),
        ("Рекомендации — «ИИ-вода»", True, " — общий план без привязки к статье закона, сумме штрафа и конкретному действию."),
    ]:
        bullet(doc, [(t, b), (rest, False)])
    body(doc, "Параллельно у холдинга есть СТ Казатомпром — стандарт безопасной эксплуатации GPS "
              "с матрицей лимитов по зонам риска × типам ТС и справочником ~400 именованных геозон. "
              "Сейчас он живёт на бумаге и в настройках Omnicomm, но не превращён в аналитику "
              "нарушений и ущерба.")

    h1(doc, 2, "Видение")
    body(doc, "Единая аналитическая платформа поверх Omnicomm — «мозг и память холдинга»: собирает "
              "телеметрию по всем ДЗО через API, нормализует в иерархию организаций "
              "(Холдинг → ДЗО → под-ДЗО → подрядчик), переводит в деньги (₸) и раздаёт каждому ДЗО "
              "свой дашборд и отчёты — без посемест­ных лицензий, с доступом строго к своей технике.")
    body(doc, "Платформа не переизобретает то, что уже считает Omnicomm (расход, заправки/сливы, "
              "моточасы, SafeDrive), и не лезет в видео-железо (контур интегратора). Её ценность — "
              "слой, которого нет ни у Omnicomm, ни у Power BI, ни у видеоинтегратора: кросс-ДЗО "
              "агрегация + экономика + рекомендации на букве закона + контроль ТО и качества данных.")

    h1(doc, 3, "Ценностное предложение")
    add_table(doc,
        ["Пласт", "Что даёт", "Кому больно сейчас"],
        [
            ["Экономика топлива", "Перерасход/экономия в ₸, нормы, рейтинги ДЗО", "нет в Power BI"],
            ["Скоростной режим на букве закона", "Превышения по геозоне × типу ТС (СТ КАП), квалификация по КоАП РК, сумма штрафов, вероятный ущерб, действие", "Power BI показывает факт, не деньги/закон"],
            ["КПД спецтехники", "Моточасы, простой под нагрузкой, ₸/моточас (буровые, погрузка)", "не считает никто"],
            ["Контроль ТО", "Авто-учёт наработки, заблаговременный алерт, подтверждение цикла", "вручную/никак"],
            ["Качество данных (Sensor Health)", "«Светофор» терминалов, пометка ТС с битыми датчиками", "аналитик ловит глазами"],
            ["Доставка", "Дашборд на ДЗО + Excel + 2 языка + мобильный вид, без лицензий", "лицензии + ручной экспорт"],
        ],
        widths=[Pt(120), Pt(250), Pt(120)])
    body(doc, "")
    body(doc, "Главный дифференциатор — рекомендации на букве закона: нарушение → статья КоАП РК → "
              "частота → сумма штрафов (МРП × ₸) → вероятный ущерб → конкретное действие (беседа с "
              "водителем; при систематике — дооснащение DSM/ADAS как инструмент продаж интегратора).",
              bold=False)

    h1(doc, 4, "Позиционирование (с кем не конкурируем)")
    for t, rest in [
        ("vs Omnicomm Online", " — мы надстройка, не замена. Omnicomm остаётся источником телеметрии и считает per-ТС; мы агрегируем по холдингу и переводим в управленческие решения."),
        ("vs Power BI", " — заменяем как витрину: дешевле (без лицензий на получателя), масштабируемо (иерархия), богаче (деньги/закон/ТО)."),
        ("vs видеоинтегратор (IT-Transport)", " — у них «глаза и нервы» (DMS/ADAS/MDVR, отдельный контур). У них нет ИИ-аналитики, скоринга и холдингового среза — это наш участок."),
    ]:
        bullet(doc, [(t, True), (rest, False)])

    h1(doc, 5, "Бизнес-модель")
    bullet(doc, [("«Аналитика как услуга» ", True), ("поверх существующей телеметрии холдинга — без продажи железа.", False)])
    bullet(doc, [("Upsell-воронка: ", True), ("платформа автоматически обосновывает дооснащение (систематические нарушения → DSM/ADAS), контроль ТО и качества данных.", False)])
    bullet(doc, [("Защита данных: ", True), ("вся инфраструктура в РК, без трансграничной передачи ПДн.", False)])

    h1(doc, 6, "Границы (scope)")
    body(doc, "Делаем:", bold=True, after=2)
    body(doc, "сбор из Omnicomm API, иерархия организаций, дашборды/отчёты на ДЗО, экономика, "
              "скоростной режим по СТ КАП, рекомендации на букве закона, контроль ТО, терминальный "
              "Sensor Health, экспорт Excel, 2 языка, мобильный вид.")
    body(doc, "Не делаем:", bold=True, after=2)
    body(doc, "видео-железо и видеоаналитику (контур интегратора); не переизобретаем расчёты "
              "Omnicomm; не лезем в профильные системы учёта Казатомпрома (доступа нет); не строим "
              "сенсор-уровневую детекцию сбоев, пока нет доступа к «Журналу» через API.")

    h1(doc, 7, "Стейкхолдеры")
    for t, rest in [
        ("Заказчик", " — руководство КАП (видит весь холдинг) + ДЗО (видят свою технику)."),
        ("Аналитик заказчика", " — ведёт текущую витрину, формулирует требования multi-tenant."),
        ("Директор по развитию Omnicomm", " — принцип «на букве закона» и upsell-логику."),
        ("Интегратор (IT-Transport)", " — телематика, монтаж, видео-контур, доступ к данным."),
    ]:
        bullet(doc, [(t, True), (rest, False)])

    h1(doc, 8, "Архитектура и текущий статус (на 22.06.2026)")
    body(doc, "Платформа переведена на современный стек, ориентированный на скорость и "
              "масштабируемость:", after=4)
    for t, rest in [
        ("Python-движки", " — нормализация, экономика, скоростной режим СТ КАП, роллапы по иерархии (переиспользуют отлаженный single-client движок)."),
        ("FastAPI-мост", " — синхронизация с Omnicomm идёт фоновой задачей с прогрессом и параллельным забором; результат кладётся в кэш-снапшот."),
        ("Кэш-снапшот", " — чтения дашборда отдаются из кэша мгновенно и не обращаются к Omnicomm на каждый запрос: интерфейс всегда быстрый."),
        ("Next.js-интерфейс", " — современный дашборд ДЗО, карта геозон, рекомендации; экспорт и два языка на дорожной карте."),
        ("Устойчивость", " — синк изолирован: его сбой или медлительность Omnicomm не роняет сервер; данные переживают перезапуск."),
    ]:
        bullet(doc, [(t, True), (rest, False)])
    body(doc, "Подключена боевая копия КАП (online.omnicomm.ru): дерево → 73 узла / ~1998 ТС, "
              "топливо и обороты идут по API; получен СТ КАП с геозонами. Платформа развёрнута на "
              "постоянном домене. Дальнейшие шаги — расширение детекции скоростного режима, "
              "Excel-экспорт, второй язык и мобильный вид.", after=4)

    # подпись-футер
    body(doc, "")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TechnoKod · аналитика автопарка как услуга поверх Omnicomm")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED; _runfont(r)

    doc.save(str(OUT))
    print("saved:", OUT, OUT.stat().st_size, "bytes")


if __name__ == "__main__":
    build()
